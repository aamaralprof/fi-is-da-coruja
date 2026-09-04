# -*- coding: utf-8 -*-
"""Gera passaportes para uma turma.

Produz duas coisas, e a diferença entre elas é o coração do sistema:

  saida/passaportes.sql   vai para o banco. Só códigos e PINs cifrados.
  saida/etiquetas.html    fica com a professora. Nomes, códigos e PINs legíveis.

O nome do aluno aparece na etiqueta e em nenhum outro lugar. Ele não entra no
SQL, não vai para a Cloudflare, não sai deste computador. O banco continua
sabendo apenas que CORUJA-7K4M leu o capítulo 4 — nunca quem é.

Uso:
    python gerar_passaportes.py --nomes turma.txt --turma "7º B"
    python gerar_passaportes.py 32 --turma "7º B"     (sem nomes)
"""

import argparse
import base64
import hashlib
import html
import io
import os
import secrets

# Sem I, L, O, 0 e 1: à mão, num papel recortado, essas letras viram outras.
ALFABETO = "ABCDEFGHJKMNPQRSTUVWXYZ23456789"
ITERACOES = 100000  # precisa ser igual ao do worker.js
AQUI = os.path.dirname(os.path.abspath(__file__))
SAIDA = os.path.join(AQUI, "saida")
SELO = os.path.join(AQUI, "..", "blog-sofia", "assets", "escritorio-do-destino-inventario.png")


def base64url(dados: bytes) -> str:
    return base64.urlsafe_b64encode(dados).decode().rstrip("=")


def novo_codigo() -> str:
    return "CORUJA-" + "".join(secrets.choice(ALFABETO) for _ in range(4))


def novo_pin() -> str:
    return "".join(secrets.choice("0123456789") for _ in range(4))


def cifrar(pin: str, sal: bytes) -> str:
    return base64url(hashlib.pbkdf2_hmac("sha256", pin.encode(), sal, ITERACOES, dklen=32))


# Cabeçalhos, rodapés e colunas que uma lista de escola costuma trazer junto
# e que não são nome de aluno.
RUIDO = (
    "nome", "aluno", "aluna", "estudante", "chamada", "ra ", "nº", "no.",
    "turma", "série", "serie", "escola", "professor", "diretoria",
    "matrícula", "matricula", "total", "relação", "relacao",
    "lista", "classe", "ensino", "fundamental", "médio", "medio",
    "secretaria", "educação", "educacao", "governo", "estado",
)


def parece_nome(texto: str) -> bool:
    """Filtro grosseiro para separar nome de aluno do resto da planilha."""
    t = " ".join(texto.split())
    if len(t) < 5 or len(t) > 70:
        return False
    if any(c.isdigit() for c in t):
        return False
    if len(t.split()) < 2:          # nome de aluno costuma ter sobrenome
        return False
    baixo = t.lower()
    if any(baixo.startswith(r) or baixo == r.strip() for r in RUIDO):
        return False
    return True


def _de_texto(caminho):
    with io.open(caminho, encoding="utf-8", errors="replace") as f:
        return [l.strip() for l in f if l.strip() and not l.lstrip().startswith("#")]


def _de_csv(caminho):
    import csv
    linhas = []
    with io.open(caminho, encoding="utf-8-sig", errors="replace", newline="") as f:
        amostra = f.read(4096); f.seek(0)
        try:
            dialeto = csv.Sniffer().sniff(amostra, delimiters=",;	")
        except Exception:
            dialeto = csv.excel
        for linha in csv.reader(f, dialeto):
            linhas.extend(c for c in linha if c and c.strip())
    return linhas


def _cabecalho_de(linha):
    return [" ".join(str(c or "").split()).lower() for c in linha]


def _coluna_por(cabecalho, *palavras):
    for i, celula in enumerate(cabecalho):
        if any(p in celula for p in palavras):
            return i
    return None


def _nomes_de_tabela(tabela):
    """Extrai a coluna de nomes de uma tabela com cabeçalho.

    Listas de escola vêm em tabela: número de chamada, nome, RA, e-mails,
    situação. Ler linha a linha grudaria tudo isso num nome só, por isso
    localizamos a coluna certa pelo cabeçalho.
    """
    if not tabela or len(tabela) < 2:
        return None
    cabecalho = _cabecalho_de(tabela[0])
    coluna_nome = _coluna_por(cabecalho, "nome", "aluno", "estudante")
    if coluna_nome is None:
        return None
    coluna_situacao = _coluna_por(cabecalho, "situa")

    nomes, dispensados = [], []
    for linha in tabela[1:]:
        if coluna_nome >= len(linha):
            continue
        nome = " ".join(str(linha[coluna_nome] or "").split())
        if not nome:
            continue
        situacao = ""
        if coluna_situacao is not None and coluna_situacao < len(linha):
            situacao = " ".join(str(linha[coluna_situacao] or "").split())
        # Quem saiu da escola não recebe passaporte.
        if situacao and situacao.strip().lower() not in ("ativo", "ativa"):
            dispensados.append((nome, situacao))
            continue
        nomes.append(nome)
    return nomes, dispensados


DISPENSADOS = []


def _de_excel(caminho):
    from openpyxl import load_workbook
    livro = load_workbook(caminho, data_only=True, read_only=True)
    tabelas = []
    for aba in livro.worksheets:
        linhas = [list(l) for l in aba.iter_rows(values_only=True)]
        # o cabeçalho pode não estar na primeira linha do arquivo
        for inicio, linha in enumerate(linhas[:15]):
            if _coluna_por(_cabecalho_de(linha), "nome", "aluno", "estudante") is not None:
                tabelas.append(linhas[inicio:]); break
    livro.close()

    for tabela in tabelas:
        achado = _nomes_de_tabela(tabela)
        if achado and achado[0]:
            DISPENSADOS.extend(achado[1])
            return ("estruturado", achado[0])

    livro = load_workbook(caminho, data_only=True, read_only=True)
    celulas = []
    for aba in livro.worksheets:
        for linha in aba.iter_rows(values_only=True):
            celulas.extend(str(v) for v in linha if v is not None)
    livro.close()
    return celulas


def _de_pdf(caminho):
    import pdfplumber
    nomes = []
    with pdfplumber.open(caminho) as pdf:
        for pagina in pdf.pages:
            for tabela in pagina.extract_tables():
                achado = _nomes_de_tabela(tabela)
                if achado:
                    nomes.extend(achado[0])
                    DISPENSADOS.extend(achado[1])
    if nomes:
        return ("estruturado", nomes)

    linhas = []
    with pdfplumber.open(caminho) as pdf:
        for pagina in pdf.pages:
            texto = pagina.extract_text() or ""
            linhas.extend(l.strip() for l in texto.splitlines() if l.strip())
    return linhas


def _de_word(caminho):
    import docx
    documento = docx.Document(caminho)
    for tabela in documento.tables:
        linhas = [[c.text for c in linha.cells] for linha in tabela.rows]
        achado = _nomes_de_tabela(linhas)
        if achado and achado[0]:
            DISPENSADOS.extend(achado[1])
            return ("estruturado", achado[0])
    pedacos = [p.text for p in documento.paragraphs]
    for tabela in documento.tables:
        for linha in tabela.rows:
            pedacos.extend(c.text for c in linha.cells)
    return [p.strip() for p in pedacos if p and p.strip()]


LEITORES = {
    ".txt": _de_texto, ".csv": _de_csv, ".tsv": _de_csv,
    ".xlsx": _de_excel, ".xlsm": _de_excel,
    ".pdf": _de_pdf, ".docx": _de_word,
}


def ler_nomes(caminho: str, filtrar=None):
    """Lê nomes de txt, csv, excel, pdf ou word.

    Em txt escrito à mão, confia no que está lá. Nos outros formatos, que vêm
    de sistemas da escola e trazem cabeçalho, número de chamada e RA no meio,
    aplica um filtro e mostra o resultado para conferência.
    """
    if not os.path.isabs(caminho):
        tentativa = os.path.join(AQUI, caminho)
        caminho = tentativa if os.path.exists(tentativa) else caminho
    if not os.path.exists(caminho):
        raise SystemExit("nao encontrei o arquivo: " + caminho)

    extensao = os.path.splitext(caminho)[1].lower()
    leitor = LEITORES.get(extensao)
    if not leitor:
        raise SystemExit(
            "nao sei ler arquivos {}. Formatos aceitos: {}".format(
                extensao or "sem extensao", ", ".join(sorted(LEITORES)))
        )

    bruto = leitor(caminho)

    # Tabela reconhecida: a coluna de nomes já veio limpa, sem chute.
    if isinstance(bruto, tuple) and bruto and bruto[0] == "estruturado":
        nomes = bruto[1]
    else:
        if filtrar is None:
            filtrar = extensao != ".txt"
        candidatos = [" ".join(c.split()) for c in bruto]
        nomes = [c for c in candidatos if parece_nome(c)] if filtrar else candidatos

    vistos, unicos = set(), []
    for n in nomes:
        chave = n.lower()
        if chave not in vistos:
            vistos.add(chave); unicos.append(n)
    return unicos


def selo_embutido() -> str:
    """O selo entra dentro do arquivo, não como link.

    Assim a folha de etiquetas pode ser movida, copiada ou aberta em outro
    computador sem perder a estampa.
    """
    try:
        from PIL import Image
        with Image.open(SELO) as im:
            im = im.convert("RGBA")
            im.thumbnail((190, 190), Image.LANCZOS)
            buffer = io.BytesIO()
            im.save(buffer, format="PNG", optimize=True)
            bruto = buffer.getvalue()
    except Exception:
        try:
            with open(SELO, "rb") as f:
                bruto = f.read()
        except Exception:
            return ""
    return "data:image/png;base64," + base64.b64encode(bruto).decode()


def gerar(nomes):
    passaportes, vistos = [], set()
    for nome in nomes:
        codigo = novo_codigo()
        while codigo in vistos:
            codigo = novo_codigo()
        vistos.add(codigo)
        pin = novo_pin()
        sal = secrets.token_bytes(16)
        passaportes.append({
            "nome": nome,
            "codigo": codigo,
            "pin": pin,
            "sal": base64url(sal),
            "hash": cifrar(pin, sal),
        })
    return passaportes


def escrever_sql(passaportes, caminho):
    """Grava um único comando, numa única linha.

    O console do D1 engasga com dezenas de comandos colados de uma vez. Um
    INSERT com muitas linhas de VALUES faz o mesmo serviço e é um comando só,
    que se seleciona, cola e executa sem susto. O arquivo não leva comentário
    nenhum, para que "selecionar tudo e colar" seja sempre seguro.
    """
    valores = ", ".join(
        "('{codigo}', '{hash}', '{sal}', datetime('now'))".format(**p)
        for p in passaportes
    )
    comando = (
        "INSERT INTO passaportes (codigo, pin_hash, pin_sal, criado_em) VALUES "
        + valores
        + " ON CONFLICT(codigo) DO NOTHING;"
    )
    quebra = chr(10)
    with io.open(caminho, "w", encoding="utf-8", newline=quebra) as f:
        f.write(comando + quebra)


ETIQUETA = """<article class="etiqueta">
  <div class="selo">{selo}</div>
  <p class="emissor">EMITIDO PELO ESCRITÓRIO DO DESTINO</p>
  <p class="nome">{nome}</p>
  <div class="credenciais">
    <span><small>CÓDIGO</small><strong>{codigo}</strong></span>
    <span><small>PIN</small><strong class="pin">{pin}</strong></span>
  </div>
  <p class="rodape">nem toda passagem aparece no mapa</p>
</article>"""


def escrever_etiquetas(passaportes, caminho, turma):
    selo = selo_embutido()
    marca = '<img src="{}" alt="">'.format(selo) if selo else "&#9670;"

    corpo = "".join(
        ETIQUETA.format(
            selo=marca,
            nome=html.escape(p["nome"]) or "&nbsp;",
            codigo=html.escape(p["codigo"]),
            pin=html.escape(p["pin"]),
        )
        for p in passaportes
    )

    titulo = "Passaportes" + (" — " + html.escape(turma) if turma else "")
    pagina = u"""<!doctype html>
<html lang="pt-BR"><head><meta charset="utf-8"><title>{titulo}</title>
<link rel="stylesheet" href="https://fonts.googleapis.com/css2?family=Cormorant+Garamond:wght@600&family=DM+Sans:wght@400;500;700&display=swap">
<style>
*{{box-sizing:border-box}}
body{{margin:0;padding:10mm;color:#2b2419;background:#fff;font-family:"DM Sans",system-ui,sans-serif}}
h1{{margin:0 0 3mm;font-family:"Cormorant Garamond",Georgia,serif;font-size:17pt;font-weight:600}}
.aviso{{max-width:190mm;margin:0 0 8mm;padding:4mm 5mm;font-size:8.5pt;line-height:1.55;
  background:#faf6ec;border-left:3px solid #b3922f}}
.folha{{display:grid;grid-template-columns:repeat(2,1fr);gap:4mm}}
.etiqueta{{position:relative;display:flex;flex-direction:column;align-items:center;
  padding:6mm 5mm 5mm;text-align:center;break-inside:avoid;
  background:#fdfaf2;border:1px dashed #a8977a;border-radius:1.5mm}}
.selo{{height:15mm;margin-bottom:1.5mm}}
.selo img{{height:15mm;width:auto;display:block}}
.emissor{{margin:0 0 2.5mm;color:#8a7442;font-size:5.6pt;font-weight:700;letter-spacing:.13em}}
.nome{{margin:0 0 3.5mm;padding-bottom:2.5mm;width:100%;
  font-family:"Cormorant Garamond",Georgia,serif;font-size:14.5pt;line-height:1.15;
  border-bottom:1px solid #d9cdb2}}
.credenciais{{display:flex;justify-content:center;gap:8mm;width:100%}}
.credenciais span{{display:flex;flex-direction:column;gap:.8mm}}
.credenciais small{{color:#96876a;font-size:5.6pt;font-weight:700;letter-spacing:.13em}}
.credenciais strong{{font-size:12pt;letter-spacing:.05em}}
.pin{{letter-spacing:.3em}}
.rodape{{margin:3.5mm 0 0;color:#a2937a;font-size:6pt;font-style:italic}}
@media print{{
  body{{padding:8mm}}
  .aviso{{background:none;border-left-color:#999}}
  .etiqueta{{background:none}}
}}
</style></head><body>
<h1>{titulo}</h1>
<p class="aviso"><strong>Esta folha é a única ligação entre um código e um aluno.</strong>
O banco de dados guarda apenas o código e o PIN cifrado — o nome não sai daqui.
Confira os nomes, recorte e entregue. <strong>Guarde uma via.</strong> Se ela se
perder, ninguém mais consegue dizer de quem é cada passaporte.</p>
<div class="folha">{corpo}</div>
</body></html>""".format(titulo=titulo, corpo=corpo)

    with io.open(caminho, "w", encoding="utf-8") as f:
        f.write(pagina)


def _borda_tracejada(celula):
    """python-docx não expõe bordas de célula; é preciso descer ao XML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    propriedades = celula._tc.get_or_add_tcPr()
    bordas = OxmlElement("w:tcBorders")
    for lado in ("top", "left", "bottom", "right"):
        elemento = OxmlElement("w:" + lado)
        elemento.set(qn("w:val"), "dashed")
        elemento.set(qn("w:sz"), "6")
        elemento.set(qn("w:color"), "A8977A")
        bordas.append(elemento)
    propriedades.append(bordas)


def escrever_docx(passaportes, caminho, turma):
    """A mesma folha de etiquetas, em Word, para ajustar antes de imprimir."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    documento = Document()
    secao = documento.sections[0]
    for lado in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(secao, lado, Cm(1.4))

    titulo = documento.add_paragraph()
    corrida = titulo.add_run("Passaportes" + (" — " + turma if turma else ""))
    corrida.font.size = Pt(16)
    corrida.font.bold = True

    aviso = documento.add_paragraph()
    a1 = aviso.add_run("Esta folha é a única ligação entre um código e um aluno. ")
    a1.font.bold = True
    a1.font.size = Pt(8.5)
    a2 = aviso.add_run(
        "O banco de dados guarda apenas o código e o PIN cifrado — o nome não sai "
        "daqui. Confira os nomes, recorte e entregue. Guarde uma via: se ela se "
        "perder, ninguém mais consegue dizer de quem é cada passaporte."
    )
    a2.font.size = Pt(8.5)

    selo = None
    try:
        from PIL import Image
        with Image.open(SELO) as im:
            im = im.convert("RGBA")
            im.thumbnail((220, 220), Image.LANCZOS)
            selo = io.BytesIO()
            im.save(selo, format="PNG")
    except Exception:
        selo = None

    colunas = 2
    linhas = (len(passaportes) + colunas - 1) // colunas
    tabela = documento.add_table(rows=linhas, cols=colunas)
    tabela.alignment = WD_TABLE_ALIGNMENT.CENTER
    largura = Cm(8.7)

    for indice, passaporte in enumerate(passaportes):
        celula = tabela.cell(indice // colunas, indice % colunas)
        celula.width = largura
        _borda_tracejada(celula)

        def paragrafo(espaco_antes=0, espaco_depois=2):
            par = celula.add_paragraph()
            par.alignment = WD_ALIGN_PARAGRAPH.CENTER
            par.paragraph_format.space_before = Pt(espaco_antes)
            par.paragraph_format.space_after = Pt(espaco_depois)
            return par

        celula.paragraphs[0].text = ""
        if selo:
            selo.seek(0)
            marca = celula.paragraphs[0]
            marca.alignment = WD_ALIGN_PARAGRAPH.CENTER
            marca.paragraph_format.space_before = Pt(6)
            marca.paragraph_format.space_after = Pt(2)
            marca.add_run().add_picture(selo, height=Cm(1.35))

        emissor = paragrafo(0, 4).add_run("EMITIDO PELO ESCRITÓRIO DO DESTINO")
        emissor.font.size = Pt(5.5)
        emissor.font.bold = True
        emissor.font.color.rgb = RGBColor(0x8A, 0x74, 0x42)

        nome = paragrafo(0, 5).add_run(passaporte["nome"])
        nome.font.size = Pt(11.5)

        credencial = paragrafo(0, 6)
        rotulo = credencial.add_run("CÓDIGO  ")
        rotulo.font.size = Pt(5.5)
        rotulo.font.bold = True
        rotulo.font.color.rgb = RGBColor(0x96, 0x87, 0x6A)
        valor = credencial.add_run(passaporte["codigo"])
        valor.font.size = Pt(11)
        valor.font.bold = True
        separador = credencial.add_run("      PIN  ")
        separador.font.size = Pt(5.5)
        separador.font.bold = True
        separador.font.color.rgb = RGBColor(0x96, 0x87, 0x6A)
        pin = credencial.add_run(passaporte["pin"])
        pin.font.size = Pt(11)
        pin.font.bold = True

    documento.save(caminho)


ENDERECO = "https://blog-da-sofia.aamaral.workers.dev"


def conferir_banco():
    """Testa se as etiquetas geradas batem com o que está no banco.

    Não há como saber isso de fora sem tentar entrar: código inexistente e PIN
    errado respondem igual, de propósito, para que ninguém descubra a lista de
    códigos válidos. Então usamos os PINs da própria folha e amostramos alguns
    passaportes. Se abrem, a leva no banco é esta.
    """
    import json
    import urllib.request
    import urllib.error

    caminho = os.path.join(SAIDA, "etiquetas.html")
    if not os.path.exists(caminho):
        print("Nao encontrei etiquetas geradas em sistema-passaporte/saida/.")
        return 1

    import re
    pagina = io.open(caminho, encoding="utf-8").read()
    pares = re.findall(
        r'<strong>(CORUJA-[A-Z0-9]{4})</strong>.*?<strong class="pin">(\d{4})</strong>',
        pagina, re.S)
    if not pares:
        print("Nao consegui ler codigos das etiquetas.")
        return 1

    total = len(pares)
    indices = sorted(set([0, total // 4, total // 2, (3 * total) // 4, total - 1]))
    print("Conferindo %d de %d passaportes contra o banco:" % (len(indices), total))
    print("")

    abriram = 0
    for i in indices:
        codigo, pin = pares[i]
        corpo = json.dumps({"codigo": codigo, "pin": pin}).encode()
        pedido = urllib.request.Request(
            ENDERECO + "/api/entrar", data=corpo,
            headers={"Content-Type": "application/json",
                     "User-Agent": "Mozilla/5.0 (conferencia de passaportes)"})
        try:
            with urllib.request.urlopen(pedido, timeout=20) as r:
                status = r.status
        except urllib.error.HTTPError as e:
            status = e.code
        except Exception as e:
            print("  %s -> nao consegui falar com o site (%s)" % (codigo, e))
            return 1
        if status == 200:
            print("  %s -> abre" % codigo); abriram += 1
        elif status == 429:
            print("  %s -> travado por tentativas, tente daqui a 15 min" % codigo)
        else:
            print("  %s -> NAO esta no banco" % codigo)

    print("")
    if abriram == len(indices):
        print("Tudo certo: as etiquetas impressas correspondem ao banco.")
        return 0
    if abriram == 0:
        print("As etiquetas NAO correspondem ao banco.")
        print("Provavelmente o banco tem uma leva antiga. No Console do D1:")
        print("    DELETE FROM passaportes;")
        print("e depois cole o conteudo de saida/passaportes.sql.")
        return 1
    print("Alguns abrem e outros nao: o banco tem levas misturadas.")
    print("O mais seguro e limpar tudo e recarregar saida/passaportes.sql.")
    return 1


def main():
    parser = argparse.ArgumentParser(description="Gera passaportes para uma turma.")
    parser.add_argument("quantidade", nargs="?", type=int, default=None,
                        help="quantos passaportes gerar, quando não houver lista de nomes")
    parser.add_argument("--nomes", default=None,
                        help="arquivo com um nome por linha; os nomes só vão para a etiqueta")
    parser.add_argument("--turma", default="", help="rótulo impresso na folha; nunca vai ao banco")
    parser.add_argument("--conferir", action="store_true",
                        help="mostra os nomes encontrados e para, sem gerar nada")
    parser.add_argument("--sem-filtro", action="store_true",
                        help="usa todas as linhas do arquivo, sem tentar separar nome de cabeçalho")
    parser.add_argument("--refazer", action="store_true",
                        help="substitui uma leva já gerada, criando códigos novos")
    parser.add_argument("--conferir-banco", action="store_true",
                        help="testa se as etiquetas geradas batem com o banco no ar")
    argumentos = parser.parse_args()

    if argumentos.conferir_banco:
        raise SystemExit(conferir_banco())

    if argumentos.nomes:
        filtrar = False if argumentos.sem_filtro else None
        nomes = ler_nomes(argumentos.nomes, filtrar=filtrar)
        if argumentos.conferir:
            print("Encontrei {} nomes em {}:".format(len(nomes), argumentos.nomes))
            print("")
            for i, nome in enumerate(nomes, 1):
                print("  {:>3}. {}".format(i, nome))
            print("")
            if DISPENSADOS:
                print("Fora da lista, por situacao na escola:")
                for nome, situacao in DISPENSADOS:
                    print("       {}  ({})".format(nome, situacao))
                print("")
            print("Se a lista estiver certa, rode de novo sem --conferir.")
            print("Se faltou ou sobrou alguem, me mostre esta saida.")
            return
        if not nomes:
            parser.error("nao encontrei nenhum nome no arquivo. Tente com --sem-filtro")
    elif argumentos.quantidade:
        if argumentos.quantidade < 1 or argumentos.quantidade > 500:
            parser.error("escolha entre 1 e 500 passaportes")
        nomes = [""] * argumentos.quantidade
    else:
        parser.error("informe uma quantidade, um arquivo com --nomes, ou use --conferir-banco")

    caminho_sql_antigo = os.path.join(SAIDA, "passaportes.sql")
    if os.path.exists(caminho_sql_antigo) and not argumentos.refazer:
        print("Ja existe uma leva de passaportes em sistema-passaporte/saida/.")
        print("")
        print("Gerar de novo cria codigos DIFERENTES. Se voce ja imprimiu as")
        print("etiquetas ou ja carregou o SQL no banco, os novos codigos nao vao")
        print("bater com nada disso.")
        print("")
        print("Se e isso mesmo que voce quer, rode de novo com --refazer.")
        print("Antes de usar a leva nova, limpe a antiga no Console do D1 com:")
        print("    DELETE FROM passaportes;")
        raise SystemExit(1)

    os.makedirs(SAIDA, exist_ok=True)
    passaportes = gerar(nomes)

    caminho_sql = os.path.join(SAIDA, "passaportes.sql")
    caminho_etiquetas = os.path.join(SAIDA, "etiquetas.html")
    caminho_word = os.path.join(SAIDA, "etiquetas.docx")
    escrever_sql(passaportes, caminho_sql)
    escrever_etiquetas(passaportes, caminho_etiquetas, argumentos.turma)

    aviso_word = None
    try:
        escrever_docx(passaportes, caminho_word, argumentos.turma)
    except Exception as erro:
        aviso_word = str(erro)

    print("{} passaportes gerados.".format(len(passaportes)))
    if aviso_word is None:
        print("  etiquetas (Word): {}".format(caminho_word))
    print("  etiquetas (web):  {}".format(caminho_etiquetas))
    print("  banco:            {}".format(caminho_sql))
    if aviso_word:
        print("")
        print("  Nao consegui gerar o Word: {}".format(aviso_word))
        print("  A versao em HTML acima imprime igual.")
    print("")
    print("As etiquetas trazem nome, codigo e PIN legiveis.")
    print("Elas nao entram no repositorio nem no banco de dados.")
    print("")
    print("ATENCAO: cada geracao cria codigos NOVOS.")
    print("Se voce ja tinha cadastrado uma leva no banco, aqueles codigos")
    print("continuam la e nao correspondem a estas etiquetas. Antes de colar")
    print("o SQL novo, limpe os antigos com:  DELETE FROM passaportes;")
    print("(so faca isso enquanto nenhum aluno tiver comecado a usar)")


if __name__ == "__main__":
    main()
